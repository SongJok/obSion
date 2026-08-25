import re
from dataclasses import dataclass
from io import BytesIO
from pathlib import PurePath

from bs4 import BeautifulSoup
from docx import Document as WordDocument
from openpyxl import load_workbook
from pypdf import PdfReader

from obsion.common.errors import ValidationError


@dataclass(frozen=True, slots=True)
class ParsedDocument:
    text: str
    parser_version: str
    metadata: dict[str, object]


def parse_document(content: bytes, media_type: str, filename: str) -> ParsedDocument:
    normalized = media_type.split(";", 1)[0].strip().lower()
    suffix = PurePath(filename).suffix.lower()
    try:
        if normalized in {"text/plain", "text/markdown", "text/x-markdown"} or suffix in {
            ".txt",
            ".md",
            ".markdown",
        }:
            return ParsedDocument(content.decode("utf-8"), "text-v1", {})
        if normalized == "text/html" or suffix in {".html", ".htm"}:
            soup = BeautifulSoup(content, "html.parser")
            for element in soup(["script", "style", "noscript"]):
                element.decompose()
            return ParsedDocument(soup.get_text("\n", strip=True), "html-bs4-v1", {})
        if normalized == "application/pdf" or suffix == ".pdf":
            reader = PdfReader(BytesIO(content))
            pages = [page.extract_text() or "" for page in reader.pages]
            return ParsedDocument("\n\n".join(pages), "pdf-pypdf-v1", {"page_count": len(pages)})
        if (
            normalized == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            or suffix == ".docx"
        ):
            document = WordDocument(BytesIO(content))
            paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
            for table in document.tables:
                paragraphs.extend("\t".join(cell.text for cell in row.cells) for row in table.rows)
            return ParsedDocument("\n".join(paragraphs), "docx-python-docx-v1", {})
        if (
            normalized == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            or suffix == ".xlsx"
        ):
            workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
            lines: list[str] = []
            for worksheet in workbook.worksheets:
                lines.append(f"# {worksheet.title}")
                for row in worksheet.iter_rows(values_only=True):
                    if any(value is not None for value in row):
                        lines.append(
                            "\t".join("" if value is None else str(value) for value in row)
                        )
            return ParsedDocument(
                "\n".join(lines),
                "xlsx-openpyxl-v1",
                {"sheet_count": len(workbook.worksheets)},
            )
    except (UnicodeDecodeError, ValueError, OSError) as exc:
        raise ValidationError(
            "document_parse_failed", "The document could not be parsed safely"
        ) from exc
    raise ValidationError(
        "document_type_unsupported",
        "The document type is not supported",
        media_type=media_type,
        filename=filename,
    )


def chunk_document(
    text: str, *, max_chars: int = 1400, overlap_chars: int = 160
) -> list[tuple[list[str], str]]:
    cleaned = text.replace("\x00", "").replace("\r\n", "\n").strip()
    if not cleaned:
        raise ValidationError("document_empty", "The document contains no extractable text")
    heading_path: list[str] = []
    chunks: list[tuple[list[str], str]] = []
    buffer = ""
    buffer_heading: list[str] = []

    def flush() -> None:
        nonlocal buffer
        content = buffer.strip()
        if content:
            chunks.append((buffer_heading.copy(), content))
        buffer = content[-overlap_chars:] if overlap_chars and content else ""

    for block in re.split(r"\n\s*\n|(?=^#{1,6}\s)", cleaned, flags=re.MULTILINE):
        block = block.strip()
        if not block:
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", block.splitlines()[0])
        if heading:
            depth = len(heading.group(1))
            heading_path[:] = heading_path[: depth - 1]
            heading_path.append(heading.group(2).strip())
        if buffer and len(buffer) + len(block) + 2 > max_chars:
            flush()
            buffer_heading = heading_path.copy()
        elif not buffer:
            buffer_heading = heading_path.copy()
        if len(block) > max_chars:
            start = 0
            while start < len(block):
                piece = block[start : start + max_chars]
                if buffer:
                    buffer = f"{buffer}\n\n{piece}"
                    flush()
                else:
                    chunks.append((heading_path.copy(), piece))
                start += max_chars - overlap_chars
            continue
        buffer = f"{buffer}\n\n{block}" if buffer else block
    if buffer.strip():
        chunks.append((buffer_heading.copy(), buffer.strip()))
    return chunks
